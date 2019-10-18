#Импортируем нужные библиотеки для работы с word и excel файлами
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
import xlrd
import xlwt




#Я заранее создал документ с названием Headers.docx, в нем я создал два стиля для
#заголовков с нужными мне шрифтами отступами выравниванием и прочим.
#Также при создании стилей я указал уровни заголовков, так как в последующем я
#планирую собрать оглавление уже в самом Word

#Чтобы использовать созданные мною в Word стили в рабочем пространстве программы
#я открываю этот документ операцией ниже. Этот же документ будет у меня заполняться.
#Названия уже созданных стилей - SchoolHeader и SectionHeader
document = Document('Headers.docx')


#Тут я устанавливаю размеры полей всего моего документа (sections - это разделы
#по нашему, и вот я циклом на все возможные разделы расширяю это форматирование,
#ну т.е. на полностью весь документ)
sections = document.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)



#Тут я добавляю несколько новых стилей для разных видов текстов. Я мог бы это сделать
#и заранее в программе Word и сохранить эти стили в документе Headers.docx
#Тогда мне не надо было бы их создавать, они уже были бы доступны для
#использования после загрузки через функцию Document() выше
style = document.styles.add_style('InfoHeader', WD_STYLE_TYPE.PARAGRAPH) #добавляю стиль, InfoHeader - это название, WD_STYLE_TYPE.PARAGRAPH - это отнесение стиля к стилю для параграфов
font = document.styles['InfoHeader'].font #ниже меняю шрифты
font.name = 'TimesNewRoman'
font.size = Pt(10)
font.bold = False
paragraph_format = document.styles['InfoHeader'].paragraph_format #ниже устанавливаю выравнивание параграфа, отступ и т.д.
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format.space_after = Pt(4)


#То же самое что и выше только для нескольких других стилей
style = document.styles.add_style('NameWork', WD_STYLE_TYPE.PARAGRAPH)
font = document.styles['NameWork'].font
font.name = 'TimesNewRoman'
font.size = Pt(11)
font.bold = True
paragraph_format = document.styles['NameWork'].paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph_format.first_line_indent = Cm(0.62)
paragraph_format.space_after = Pt(0)
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE



style = document.styles.add_style('Authors', WD_STYLE_TYPE.PARAGRAPH)
font = document.styles['Authors'].font
font.name = 'TimesNewRoman'
font.size = Pt(11)
font.bold = False
font.italic = True
paragraph_format = document.styles['Authors'].paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


#Это полезный код - выводит на экран названия всех доступных в вашем рабочем документе стилей
#Если сомневаетесь в правильности вышенаписанных штук, запустите эти три строки
#и проверьте сами себя
paragraph_styles = [s for s in document.styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles:
    print(style.name)



#Ниже подгружаю данные из заранее подготовленных таблиц excel для заполнения моего документа
#По сути я открываю в переменную Schools мой эксель
#Создаю пустой лист, в который буду далее загонять значения из ячеек моего ексель файла
#Далее запускаю цикл, где каждую строку эксель файла записываю в лист как новый элемент
#На выходе получаем лист, элементы которого - это листы, элементы которых -
#содержимое ячеек одной строки. Если непонятно, можно print'ом повыводить разные эелементы
Schools = xlrd.open_workbook('Schools.xlsx')
SchoolsList = list()
for rownum in range(Schools.sheet_by_index(0).nrows):
    row = Schools.sheet_by_index(0).row_values(rownum)
    SchoolsList.append(row)


Sections = xlrd.open_workbook('Sections.xlsx')
SectionsList = list()
for rownum in range(Sections.sheet_by_index(0).nrows):
    row = Sections.sheet_by_index(0).row_values(rownum)
    SectionsList.append(row)


Works = xlrd.open_workbook('Works.xlsx')
WorksList = list()
for rownum in range(Works.sheet_by_index(0).nrows):
    row = Works.sheet_by_index(0).row_values(rownum)
    WorksList.append(row)



#Тут мы начинаем заполнение документа с применением нужных стилей. У меня три уровня
#данных:
#Есть общая группа "Физтех-школа"
#Внутри каждой физтех-школы есть научные секции
#В каждой научной секции есть некоторый набор работ
#Ну далее просто набор циклов, которые поочереди заполняют
#название физтех-школы -- данные секции -- работы этой секции
#                      -- данные следующей секции -- работы этой секции
#                      -- ну и так по всем секциям этой школы
#ну и потом переходим на следующую школы и заново цикл с секциями.
#Тут я еще добавляю где-то разделительные линии для красоты в виде рисунков
#Ну и стили - это сугубо мое, вы можете играться с оформлением как угодно
for i in range(len(SchoolsList)):
    document.add_paragraph(SchoolsList[i][0], style = 'SchoolHeader')
    document.add_paragraph('')
    for j in range(len(SectionsList)):
        if SectionsList[j][0] == SchoolsList[i][0]:
            document.add_picture('line1.jpg')
            document.add_paragraph(SectionsList[j][1], style = 'SectionHeader')
            document.add_picture('line1.jpg')
            document.add_paragraph('Председатель: ' + SectionsList[j][2], style = 'InfoHeader')
            document.add_paragraph('Зам. председателя: ' + SectionsList[j][3], style = 'InfoHeader')
            document.add_paragraph('Секретарь: ' + SectionsList[j][4], style = 'InfoHeader')
            document.add_picture('line2.jpg')
            document.add_paragraph('Дата: ' + str(SectionsList[j][5]) + '     Время: ' + str(SectionsList[j][6]), style = 'InfoHeader')
            document.add_paragraph('Место: '+ SectionsList[j][7], style = 'InfoHeader')
            document.add_paragraph('')
            document.add_paragraph('')
            for k in range(len(WorksList)):
                if WorksList[k][2] == SectionsList[j][1]:
                    document.add_paragraph(WorksList[k][0], style = 'NameWork')
                    document.add_paragraph(WorksList[k][1], style = 'Authors')
    document.add_page_break() #добавляем разрыв страницы


document.save('Final.docx') #сохраняем наш получившийся документ
